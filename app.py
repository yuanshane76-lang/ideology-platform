# app.py

from flask import Flask, render_template, request, jsonify, Response, stream_with_context
from src.service import chat_service_stream
from src.conversation import conversation_store
from src.ppt import PPTAgent
from src.config import settings
from src.debate import SAMPLE_TOPICS, ANTAGONIST_TYPES, stream_debate_events
from src.debate.service import create_session, get_session, delete_session, stream_single_round, stream_judge_summary
from src.debate.topic_agent import TopicAnalysisAgent
import json
import uuid

app = Flask(__name__)

ppt_agent = PPTAgent()
topic_agent = TopicAnalysisAgent()

# 页面路由
@app.route('/')
def index():
    return render_template('home.html')

@app.route('/old')
def old_index():
    return render_template('index.html')

@app.route('/home')
def home_page():
    return render_template('home.html')

@app.route('/ppt')
def ppt_page():
    return render_template('ppt.html')

@app.route('/debate')
def debate_page():
    return render_template('debate.html')

@app.route('/ppt/preview')
def ppt_preview():
    return render_template('ppt_preview.html')

# API路由：处理对话
@app.route('/api/chat', methods=['POST'])
def chat_endpoint():
    data = request.json
    user_query = data.get('query')
    conversation_id = data.get('conversation_id') # 前端如果传了就用，没传就None

    if not user_query:
        return jsonify({"error": "Empty query"}), 400

    # 使用 Flask 的 Response + stream_with_context 包装生成器
    return Response(
        stream_with_context(chat_service_stream(conversation_id, user_query)),
        mimetype='text/event-stream'
    )

# API路由：获取历史会话列表
@app.route('/api/history', methods=['GET'])
def history_endpoint():
    # 从 Store 获取真实的历史列表
    conversations = conversation_store.get_all_conversations()
    # 按更新时间倒序排列
    sorted_convs = sorted(conversations, key=lambda c: c.updated_at, reverse=True)
    
    history_list = []
    for conv in sorted_convs:
        # ========================================================
        # 【核心修改点】：优先读取数据库里的 title 字段
        # ========================================================
        # 1. 先尝试取数据库里的智能标题
        if hasattr(conv, 'title') and conv.title and conv.title.strip():
            title = conv.title
        else:
            # 2. 如果没有智能标题，再降级为“截取前15个字”
            title = "新对话"
            if conv.messages:
                first_msg = next((m for m in conv.messages if m.role == 'user'), None)
                if first_msg:
                    title = first_msg.content[:15] + "..." if len(first_msg.content) > 15 else first_msg.content
        
        history_list.append({
            "id": conv.conversation_id,
            "title": title,
            "date": conv.updated_at.strftime("%Y-%m-%d %H:%M")
        })

    return jsonify(history_list)

# API路由：获取单个会话的详细消息记录
@app.route('/api/history/<conversation_id>', methods=['GET'])
def conversation_detail(conversation_id):
    conv = conversation_store.get_conversation(conversation_id)
    if not conv:
        return jsonify({"error": "Conversation not found"}), 404
        
    messages = []
    for msg in conv.messages:
        messages.append({
            "role": msg.role,
            "content": msg.content,
            "timestamp": msg.timestamp.isoformat() if hasattr(msg.timestamp, 'isoformat') else str(msg.timestamp)
        })
    
    return jsonify(messages)

@app.route('/api/history/<conversation_id>', methods=['DELETE'])
def delete_conversation(conversation_id):
    success = conversation_store.delete_conversation(conversation_id)
    if success:
        return jsonify({"success": True, "message": "对话已删除"})
    else:
        return jsonify({"error": "Conversation not found"}), 404


@app.route('/api/debate/topics', methods=['GET'])
def get_debate_topics():
    return jsonify([
        {
            "title": topic.title,
            "description": topic.description,
            "difficulty": topic.difficulty,
            "tags": topic.tags,
        }
        for topic in SAMPLE_TOPICS
    ])


@app.route('/api/debate/antagonist-types', methods=['GET'])
def get_debate_antagonist_types():
    return jsonify(ANTAGONIST_TYPES)


@app.route('/api/debate/analyze', methods=['POST'])
def analyze_topic():
    """分析辩题，返回立场分析和理论模块"""
    data = request.json or {}
    topic = (data.get('topic') or '').strip()
    description = (data.get('description') or '').strip()
    
    if not topic:
        return jsonify({"error": "缺少辩题"}), 400
    
    try:
        analysis = topic_agent.analyze(topic, description)
        
        return jsonify({
            "session_id": str(uuid.uuid4()),
            "topic_analysis": {
                "topic": analysis.topic,
                "pro_position": analysis.pro_position,
                "con_position": analysis.con_position,
                "marxism_side": analysis.marxism_side,
                "marxism_reason": analysis.marxism_reason,
                "core_concepts": analysis.core_concepts,
                "debate_focus": analysis.debate_focus,
                "involves_marxism_stance": analysis.involves_marxism_stance,
                "stance_type": analysis.stance_type.value,
                "theory_modules": analysis.theory_modules
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/debate/stream', methods=['POST'])
def stream_debate():
    data = request.json or {}
    title = (data.get('title') or '').strip()
    description = (data.get('description') or '').strip()
    antagonist_type = (data.get('antagonist_type') or '反方').strip()

    try:
        rounds = int(data.get('rounds', 10))
    except (TypeError, ValueError):
        rounds = 10

    if not title:
        return jsonify({"error": "缺少辩题标题"}), 400

    rounds = max(1, min(10, rounds))

    def generate():
        try:
            for event in stream_debate_events(
                topic=title,
                description=description,
                antagonist_type=antagonist_type,
                rounds=rounds,
            ):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        except Exception as e:
            message = str(e)
            if "DataInspectionFailed" in message or "inappropriate content" in message:
                payload = {
                    "type": "error",
                    "message": "内容生成被安全审核拦截，建议使用更学术化、更中性的辩题表述后重试。",
                    "error_type": "content_moderation",
                }
            else:
                payload = {
                    "type": "error",
                    "message": message,
                    "error_type": "internal_error",
                }
            yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no',
        },
    )


@app.route('/api/debate/start', methods=['POST'])
def start_step_debate():
    data = request.json or {}
    title = (data.get('title') or '').strip()
    description = (data.get('description') or '').strip()
    
    try:
        max_rounds = int(data.get('max_rounds', 10))
    except (TypeError, ValueError):
        max_rounds = 10
    
    max_rounds = max(1, min(10, max_rounds))
    
    if not title:
        return jsonify({"error": "缺少辩题标题"}), 400
    
    session_id = str(uuid.uuid4())
    session = create_session(session_id, title, description, max_rounds)
    
    return jsonify({
        "session_id": session_id,
        "topic": title,
        "description": description,
        "max_rounds": max_rounds,
        "current_round": 0,
        "status": "initialized",
    })


@app.route('/api/debate/next', methods=['POST'])
def next_debate_round():
    data = request.json or {}
    session_id = data.get('session_id')
    
    if not session_id:
        return jsonify({"error": "缺少 session_id"}), 400
    
    session = get_session(session_id)
    if not session:
        return jsonify({"error": "会话不存在"}), 404
    
    if session.status == "completed":
        return jsonify({"error": "辩论已结束"}), 400
    
    if session.current_round >= session.max_rounds:
        return jsonify({"error": "已达最大轮次"}), 400
    
    def generate():
        try:
            for event in stream_single_round(session):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        except Exception as e:
            message = str(e)
            if "DataInspectionFailed" in message or "inappropriate content" in message:
                payload = {
                    "type": "error",
                    "message": "内容生成被安全审核拦截，建议使用更学术化、更中性的辩题表述后重试。",
                    "error_type": "content_moderation",
                }
            else:
                payload = {
                    "type": "error",
                    "message": message,
                    "error_type": "internal_error",
                }
            yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no',
        },
    )


@app.route('/api/debate/judge', methods=['POST'])
def judge_debate():
    data = request.json or {}
    session_id = data.get('session_id')
    
    if not session_id:
        return jsonify({"error": "缺少 session_id"}), 400
    
    session = get_session(session_id)
    if not session:
        return jsonify({"error": "会话不存在"}), 404
    
    if len(session.protagonist_messages) == 0:
        return jsonify({"error": "尚未开始辩论"}), 400
    
    if session.status == "completed":
        return jsonify({"error": "辩论已结束"}), 400
    
    def generate():
        try:
            for event in stream_judge_summary(session):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        except Exception as e:
            message = str(e)
            payload = {
                "type": "error",
                "message": message,
                "error_type": "internal_error",
            }
            yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no',
        },
    )


@app.route('/api/debate/session/<session_id>', methods=['GET'])
def get_debate_session(session_id):
    session = get_session(session_id)
    if not session:
        return jsonify({"error": "会话不存在"}), 404
    
    return jsonify(session.to_dict())


@app.route('/api/debate/session/<session_id>', methods=['DELETE'])
def delete_debate_session_route(session_id):
    delete_session(session_id)
    return jsonify({"success": True, "message": "会话已删除"})


@app.route('/api/reference/download', methods=['POST'])
def download_reference():
    """下载引用原文，优先 docx 格式，降级为 txt"""
    import re
    from urllib.parse import quote

    data = request.json
    ref_type = data.get('type', '')
    full_content = data.get('full_content', '')

    if ref_type == 'moment':
        title = data.get('title', '时政新闻')
        date = data.get('date', '')
        source = data.get('source', '')
        header_lines = [f"标题：{title}", f"日期：{date}", f"来源：{source}"]
        filename_base = title[:20]
    else:
        source = data.get('source', '理论文献')
        source = re.sub(r'思想道德与法治.*', '思想道德与法治', source)
        source = re.sub(r'毛泽东思想和中国特色社会主义理论体系概论.*', '毛泽东思想和中国特色社会主义理论体系概论', source)
        source = re.sub(r'新时代中国特色社会主义思想概论.*', '新时代中国特色社会主义思想概论', source)
        chapter = data.get('chapter', '')
        section = data.get('section', '')
        subsection = data.get('subsection', '')
        subsubsection = data.get('subsubsection', '')
        location = ' / '.join(filter(None, [chapter, section, subsection, subsubsection]))
        header_lines = [f"来源：{source}", f"章节：{location}"]
        filename_base = f"{source[:15]}_{location[:15]}".replace('/', '_')

    filename_base = re.sub(r'[\\/:*?"\u003c\u003e|\s]', '_', filename_base).strip('_')

    # full_content 已经在 reference_composer 里清理过了，直接使用
    cleaned_content = full_content

    # 尝试生成 docx
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(11)

        # 元信息标题行（加粗）
        for line in header_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(11)

        doc.add_paragraph()  # 空行

        # 正文段落
        for para in cleaned_content.split('\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                for run in p.runs:
                    run.font.size = Pt(11)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = filename_base + '.docx'
        filename_encoded = quote(filename.encode('utf-8'))
        return Response(
            buf.read(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f"attachment; filename=\"reference.docx\"; filename*=UTF-8''{filename_encoded}"
            }
        )

    except ImportError:
        # python-docx 未安装，降级为 txt
        print("[Download] python-docx not installed, falling back to txt")
        text = '\n'.join(header_lines) + '\n\n' + cleaned_content
        filename = filename_base + '.txt'
        filename_encoded = quote(filename.encode('utf-8'))
        return Response(
            text.encode('utf-8'),
            mimetype='text/plain; charset=utf-8',
            headers={
                'Content-Disposition': f"attachment; filename=\"reference.txt\"; filename*=UTF-8''{filename_encoded}",
                'Content-Type': 'text/plain; charset=utf-8'
            }
        )

@app.route('/api/ppt/templates', methods=['GET'])
def get_ppt_templates():
    """获取PPT模板列表"""
    return jsonify(ppt_agent.get_template_list())

@app.route('/api/ppt/outline', methods=['POST'])
def generate_ppt_outline():
    """生成PPT大纲（旧接口，保留兼容）"""
    data = request.json
    query = data.get('query')
    
    if not query:
        return jsonify({"success": False, "error": "缺少查询内容"}), 400
    
    result = ppt_agent.generate_outline(query)
    return jsonify(result)

@app.route('/api/ppt/outline/generate', methods=['POST'])
def generate_outline_new():
    """生成PPT大纲（新工作流 - 并行生成6章）"""
    data = request.json
    query = data.get('query')
    
    if not query:
        return jsonify({"success": False, "error": "缺少查询内容"}), 400
    
    result = ppt_agent.generate_outline(query)
    return jsonify(result)

@app.route('/api/ppt/outline/stream', methods=['POST'])
def generate_outline_stream():
    """流式生成PPT大纲"""
    import asyncio
    from flask import Response
    
    data = request.json
    query = data.get('query')
    
    if not query:
        return jsonify({"success": False, "error": "缺少查询内容"}), 400
    
    def generate():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            async_gen = ppt_agent.generate_outline_stream(query)
            
            while True:
                try:
                    event = loop.run_until_complete(async_gen.__anext__())
                    yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                except StopAsyncIteration:
                    break
        finally:
            loop.close()
    
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no'
        }
    )

@app.route('/api/ppt/generate/from-outline', methods=['POST'])
def generate_ppt_from_outline():
    """从大纲生成PPT"""
    data = request.json
    query = data.get('query')
    outline = data.get('outline')
    template_id = data.get('template_id')
    use_ai_background = data.get('use_ai_background', True)  # 默认启用AI背景
    
    if not query or not outline:
        return jsonify({"success": False, "error": "缺少必要参数"}), 400
    
    result = ppt_agent.generate_ppt_from_outline(
        query=query,
        outline=outline,
        template_name=template_id,
        use_ai_background=use_ai_background
    )
    
    return jsonify(result)

@app.route('/api/ppt/download/<ppt_id>')
def download_ppt(ppt_id):
    """下载生成的PPT文件"""
    from flask import send_file
    
    info = ppt_agent.get_ppt_file(ppt_id)
    if not info:
        return jsonify({"success": False, "error": "文件不存在"}), 404
    
    return send_file(
        info["filepath"],
        as_attachment=True,
        download_name=info["filename"]
    )


@app.route('/api/ppt/detail/<ppt_id>', methods=['GET'])
def get_ppt_detail(ppt_id):
    """获取PPT详情（用于预览）"""
    info = ppt_agent.get_ppt_file(ppt_id)
    
    if info:
        return jsonify({
            "success": True,
            "outline": info.get("outline"),
            "title": info.get("outline", {}).get("title", "PPT")
        })
    else:
        return jsonify({
            "success": False,
            "error": "PPT 文件不存在"
        })


@app.route('/api/ppt/html/themes', methods=['GET'])
def get_html_themes():
    """获取HTML主题列表"""
    return jsonify(ppt_agent.get_html_themes())


@app.route('/api/ppt/html/generate', methods=['POST'])
def generate_html_slides():
    """根据大纲生成HTML幻灯片"""
    data = request.json
    outline = data.get('outline')
    theme_name = data.get('theme', 'party_red')
    
    if not outline:
        return jsonify({"success": False, "error": "缺少大纲数据"}), 400
    
    result = ppt_agent.generate_html_slides(outline, theme_name)
    return jsonify(result)


@app.route('/api/ppt/html/slide/<session_id>/<int:slide_index>', methods=['GET'])
def get_html_slide(session_id, slide_index):
    """获取单个HTML幻灯片内容"""
    result = ppt_agent.get_html_slide(session_id, slide_index)
    return jsonify(result)


@app.route('/api/ppt/html/convert/<session_id>', methods=['POST'])
def convert_html_to_ppt(session_id):
    """将HTML幻灯片转换为PPT"""
    result = ppt_agent.convert_html_to_ppt(session_id)
    return jsonify(result)


@app.route('/api/ppt/html/generate/stream', methods=['POST'])
def generate_html_stream():
    """流式生成HTML幻灯片（SSE实时预览）"""
    import uuid
    
    data = request.json
    outline = data.get('outline')
    theme_name = data.get('theme', 'party_red')
    
    if not outline:
        return jsonify({"success": False, "error": "缺少大纲数据"}), 400
    
    def generate():
        import logging
        logger = logging.getLogger(__name__)
        
        session_id = str(uuid.uuid4())[:8]
        print(f"[SSE] Starting stream generation, session_id: {session_id}")
        logger.info(f"[SSE] Starting stream generation, session_id: {session_id}")
        
        slide_data_list = ppt_agent.html_generator.get_slide_data_list(outline)
        total = len(slide_data_list)
        print(f"[SSE] Total slides to generate: {total}")
        logger.info(f"[SSE] Total slides to generate: {total}")
        
        slides = []
        for i, slide_data in enumerate(slide_data_list):
            try:
                print(f"[SSE] Generating slide {i+1}/{total}")
                logger.info(f"[SSE] Generating slide {i+1}/{total}")
                html = ppt_agent.html_generator.generate_single_slide_html(
                    slide_data, theme_name, i, total
                )
                
                slide_info = {
                    "index": i,
                    "title": slide_data.get("title", f"幻灯片 {i+1}"),
                    "html": html
                }
                slides.append(slide_info)
                
                event_data = f"event: slide_ready\ndata: {json.dumps(slide_info, ensure_ascii=False)}\n\n"
                print(f"[SSE] Sending slide_ready event for slide {i+1}")
                logger.info(f"[SSE] Sending slide_ready event for slide {i+1}")
                yield event_data
                
            except Exception as e:
                print(f"[SSE] Error generating slide {i+1}: {e}")
                logger.error(f"[SSE] Error generating slide {i+1}: {e}")
                error_info = {"index": i, "error": str(e)}
                yield f"event: slide_error\ndata: {json.dumps(error_info, ensure_ascii=False)}\n\n"
        
        print(f"[SSE] All slides generated, saving to cache with session_id: {session_id}")
        logger.info(f"[SSE] All slides generated, saving to cache")
        cache_data = {
            "slides": slides,
            "outline": outline,
            "theme": theme_name
        }
        ppt_agent.html_slides_cache[session_id] = cache_data
        ppt_agent._save_session_to_cache(session_id, cache_data)
        
        done_info = {"session_id": session_id, "total": total}
        print(f"[SSE] Sending done event with session_id: {session_id}")
        logger.info(f"[SSE] Sending done event with session_id: {session_id}")
        yield f"event: done\ndata: {json.dumps(done_info, ensure_ascii=False)}\n\n"
        print(f"[SSE] Stream generation completed")
        logger.info(f"[SSE] Stream generation completed")
    
    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no'
        }
    )

if __name__ == '__main__':
    print("🚀 Server starting on http://0.0.0.0:6006")
    app.run(host='0.0.0.0', port=6006, debug=True, use_reloader=False)